import os
import zipfile
import xml.etree.ElementTree as ET
import requests

import win32com.client as win32

from PIL import Image
from difflib import SequenceMatcher
from docx import Document
from yattag import Doc
from classes import folder
from xml.dom import minidom

##############################################################################################
# Recursive function for directory search                                                    #
# folderList - list of objects that contains folder name and their absolute path             #
# absPath - directory where all mm folders are located                                       #
# serverID - server id number of a folder program is currently in                            #
# level - SSRS/SURS/SSTS/SUTS                                                                #
# locID - location id number for different areas of the website where data can be displayed  #
# e.g. description/precondition/steps boxes of files on Matrix Website                       #
# prefix - prefix of old naming convention from Tracer                                       #
# project - project name (assigned by user)                                                  #
# credentials - user credentials to log in to the Matrix server                              #
# folderID - number that represents id of a contents element in a folder                     #
# url - url assigned by user to connect to the right website                                 #
##############################################################################################

def directorySearch(folderList, absPath, serverID, level, locID, prefix, project, credentials, folderID, url):
    # Goes through all objects in a list
    for ele in folderList:
        directory = os.listdir(ele.path)
        name = ele.name

        foldDes = ""
        # Checking if there are any header files
        for file in directory:
            if(file.startswith("header")):
                file_name = ele.path + "\\" + file
                foldDes = wordReadText(file_name)

        # Gather data of the new folder
        data = {
            "reason": "test",
            "label": name,
            # Indicates where to create a new folder - i.e. inside which folder
            "parent": "F-{}-{}".format(level, serverID),
            # Add some information to contents part of the folder if it exists
            "fx{}".format(folderID): foldDes
        }

        # Upload the new folder
        resp = requests.post("http://{}{}/rest/1/{}/folder".format(credentials, url, project), data=data)

        # Get newly created folder ID (used when creating it's subfolders/elements inside of it)
        index = resp.json()["serial"]

        # Get the structure inside that newly created folder
        folders = structureFolder(ele.path + "\\structure.xml")

        # Checks if there are any folders - main part of importing done here
        if not folders:

            # List that is going to contain elements inside a folder in the right order
            # Reads structure.xml file to get that information
            order = []
            for file in directory:
                if (file == "structure.xml"):
                    order = structureXML(ele.path + "\\" + file)



            # If the folder list is not empty
            if (order != []):

                # Loop through the order list
                for file in order:
                    # Save absolute directory of the current file in the order list
                    file_name = ele.path + "\\" + file + ".docx"

                    # Open that file with python-docx library
                    document = Document(file_name)

                    # Extract word document properties
                    # Files with hash names have real names saved under title property
                    properties = document.core_properties
                    title = properties.title

                    # If the the title property is not assigned to anything, save the name of the file as a title
                    if (properties.title == "[Blank Template]"):
                        title = ele.name

                    # Get costum property of the document - tracer ID number
                    label = costumProperty(file_name)

                    # Build whole title with tracer identification number before the name of the document
                    title = prefix + label + " " + title


                    description = ""

                    # Standard upload pattern (commented):

                    # One way for requirements
                    if (level == "SSRS" or level == "SURS"):

                        # Get all info for a requirement and upload it using one function
                        requirements(file_name, absPath, document, title, level, index, locID, project, credentials, url)

                    # Otherwise upload a test
                    else:
                        # Check if test format is standard
                        if (len(document.tables) == 1 and len(document.tables[0].rows)==4 and (document.tables[0].rows[0].cells[0].text=="Description")):
                            print(file_name)
                            tests(file_name, document, title, level, index, locID, project, credentials, url)


                        # If there is only text in a test document (simple folder information)
                        elif(document.tables==[] and document.paragraphs!=[]):
                            everything = wordReadText(file_name)
                            filedata = {
                                'title': title,
                                'folder': 'F-{}-{}'.format(level, index),
                                'reason': 'test',
                                'fx{}'.format(locID): everything
                            }

                            resp = requests.post(
                                'http://{}{}/rest/1/{}/item/'.format(
                                    credentials, url, project), data=filedata)

                        # Anything more complex uses local API to find order of elements and upload them to the website
                        else:
                            #zipping a file - need to access .xml files
                            z = zipfile.ZipFile(file_name)

                            # Getting the order of elements in the correct format
                            Order = elementOrder(z)
                            Order = modifyOrder(Order)

                            # Add a warning sign because functions above are not 100% accurate
                            everything = "<p><b>Attention - Double Check!</b></p>"

                            # Add all information of the document to a string
                            everything = everything + fullDoc(file_name, Order)

                            # Format the data in the correct JSON format
                            filedata = {
                                'title': title,
                                'folder': 'F-{}-{}'.format(level, index),
                                'reason': 'test',
                                'fx{}'.format(locID): everything
                            }

                            # Upload the data to the website
                            resp = requests.post(
                                'http://{}{}/rest/1/{}/item/'.format(
                                    credentials, url, project), data=filedata)

                # do the same but for not hash names - has to go into folder
                for notHash in directory:
                    if (notHash.startswith("header")):
                        file_name = ele.path + "\\" + notHash
                        print(file_name)
                        document = Document(file_name)
                        properties = document.core_properties
                        title = properties.title

                        if (properties.title == "[Blank Template]"):
                            title = ele.name

                        label = costumProperty(file_name)

                        # If the file doesn't have a Tracer id (found with costumProperty() function),
                        # the document is a header file - id equals to "ZZZZ"
                        if (label == "ZZZZ"):
                            # Same procedure of uploading as above
                            if (level == "SSRS" or level == "SURS"):

                                # Changing folder
                                description = wordReadText(file_name)

                                filedata = {
                                    "title": name,
                                    "reason": "test",
                                    "fx{}".format(folderID): description
                                }
                                resp = requests.put(
                                    'http://{}{}/rest/1/{}/item/F-{}-{}'.format(credentials, url,
                                    project, level, index), data=filedata)

                            else:

                                if (len(document.tables) == 1 and len(document.tables[0].rows)==4 and (document.tables[0].rows[0].cells[0].text=="Description")):

                                    tests(file_name, document, title, level, index, locID, project, credentials, url)

                                # If there is only text in a document
                                elif(document.tables == [] and document.paragraphs != []):
                                    everything = wordReadText(file_name)
                                    filedata = {
                                        'title': name,
                                        'reason': 'test',
                                        'fx{}'.format(folderID): everything
                                    }

                                    resp = requests.put(
                                        'http://{}{}/rest/1/{}/item/F-{}-{}'.format(
                                            credentials, url, project, level, index), data=filedata)

                                # Non-generic tests
                                else:
                                    z = zipfile.ZipFile(file_name)

                                    Order = elementOrder(z)
                                    Order = modifyOrder(Order)

                                    everything = "<p><b>Attention - Double Check!</b></p>"
                                    everything = everything + fullDoc(file_name, Order)

                                    filedata = {
                                        'title': name,
                                        'reason': 'test',
                                        'fx{}'.format(folderID): everything
                                    }

                                    resp = requests.put(
                                        'http://{}{}/rest/1/{}/item/F-{}-{}'.format(
                                            credentials, url, project, level, index), data=filedata)

            #no hash names - just header.docx
            else:

                for file in directory:

                    # Try to to find where ~names come from - skip them
                    # "~" indicates the file is open at the time of upload or the file is corrupted
                    # What ever is the cause of it, it stops our program from uploading
                    # --> python-docx library can't open them - breaks the program
                    if (file.startswith("~") != True):

                        if (file.endswith(".docx")):

                            # Standard upload pattern (commented above when it appears for the first time)
                            file_name = ele.path + "\\" + file
                            print(file_name)
                            document = Document(file_name)
                            properties = document.core_properties
                            title = properties.title

                            if (properties.title == "[Blank Template]"):
                                title = ele.name

                            label = costumProperty(file_name)
                            if (label == "ZZZZ"):
                                if (level == "SSRS" or level == "SURS"):
                                    #Changing folder
                                    #description = wordReadText(file_name)
                                    requirementsFolder(file_name, absPath, document, title, level, index, project, credentials, folderID, url)
                                elif(document.tables == [] and document.paragraphs != []):
                                    everything = wordReadText(file_name)

                                    filedata = {
                                        'title': name,
                                        'reason': 'test',
                                        'fx{}'.format(folderID): everything
                                    }

                                    resp = requests.put(
                                        'http://{}{}/rest/1/{}/item/F-{}-{}'.format(
                                            credentials, url, project, level, index), data=filedata)

                                else:

                                    z = zipfile.ZipFile(file_name)

                                    Order = elementOrder(z)
                                    Order = modifyOrder(Order)

                                    everything = "<p><b>Attention - Double Check!</b></p>"
                                    everything = everything  + fullDoc(file_name, Order)

                                    filedata = {
                                        'title': name,
                                        'reason': 'test',
                                        'fx{}'.format(folderID): everything
                                    }

                                    resp = requests.put(
                                        'http://{}{}/rest/1/{}/item/F-{}-{}'.format(
                                            credentials, url, project, level, index), data=filedata)

                            # Might be useless
                            # Just trying to catch everything in case some heard file doesn't have a
                            # costum property ID that equals to "ZZZZ"
                            else:
                                print("Header without ID ZZZZ!")
                                title = prefix + label + " " + title

                                # Standard upload pattern

                                # One way for requirements
                                if (level == "SSRS" or level == "SURS"):

                                    ############################################################
                                    # Check if this function should be requirementsFolders     #
                                    # requirements - uploads a new file                        #
                                    # requirementsFolders - uploads to already existing folder #
                                    ############################################################
                                    requirements(file_name, absPath, document, title, level, index, locID, project, credentials, url)

                                elif (document.tables == [] and document.paragraphs != []):
                                    everything = wordReadText(file_name)

                                    filedata = {
                                        'title': name,
                                        'reason': 'test',
                                        'fx{}'.format(folderID): everything
                                    }

                                    resp = requests.put(
                                        'http://{}{}/rest/1/{}/item/F-{}-{}'.format(
                                            credentials, url, project, level, index), data=filedata)


                                else:


                                    z = zipfile.ZipFile(file_name)

                                    Order = elementOrder(z)
                                    Order = modifyOrder(Order)

                                    everything = "<p><b>Attention - Double Check!</b></p>"
                                    everything = everything + fullDoc(file_name, Order)

                                    filedata = {
                                        'title': name,
                                        'reason': 'test',
                                        'fx{}'.format(folderID): everything
                                    }

                                    resp = requests.put(
                                        'http://{}{}}/rest/1/{}/item/F-{}-{}'.format(
                                            credentials, url, project, level, index), data=filedata)

        # If there are more folders check if there are any files with hash names and then do the same as above
        else:

            # Same procedure as when there are no Folder present
            # Get new subfolders
            order = []
            for file in directory:
                if (file == "structure.xml"):
                    order = structureXML(ele.path + "\\" + file)

            # If the folder list is not empty
            if (order != []):

                for file in directory:

                    # Checking if there is a header files
                    # If it exists upload it to the folder
                    if (file.startswith("header")):
                        file_name = ele.path + "\\" + file
                        document = Document(file_name)
                        properties = document.core_properties
                        title = properties.title
                        if (properties.title == "[Blank Template]"):
                            title = ele.name
                        requirementsFolder(file_name, absPath, document, title, level, index, project, credentials, folderID, url)

                #file is just name of the item without extension
                for file in order:

                    #Getting all information for the upload (same as above - before standard upload)
                    file_name = ele.path + "\\" + file + ".docx"

                    document = Document(file_name)
                    properties = document.core_properties
                    title = properties.title

                    if (properties.title == "[Blank Template]"):
                        title = ele.name

                    label = costumProperty(file_name)

                    title = prefix + label + " " + title


                    description = ""

                    # Again standard upload procedure:

                    # One way for requirements
                    if (level == "SSRS" or level == "SURS"):
                        requirements(file_name, absPath, document, title, level, index, locID, project, credentials, url)

                    #one for tests
                    else:
                        if (len(document.tables) == 1 and len(document.tables[0].rows)==4 and (document.tables[0].rows[0].cells[0].text=="Description")):
                            print(file_name)
                            tests(file_name, document, title, level, index, locID, project, credentials, url)

                        #Only text in tests:
                        elif(document.tables==[] and document.paragraphs != []):
                            everything = wordReadText(file_name)

                            filedata = {
                                'title': title,
                                'folder': 'F-{}-{}'.format(level, index),
                                'reason': 'test',
                                'fx{}'.format(locID): everything
                            }

                            resp = requests.post(
                                'http://{}{}/rest/1/{}/item/'.format(credentials, url,
                                                                     project), data=filedata)

                        # Non-generic tests
                        else:
                            z = zipfile.ZipFile(file_name)

                            Order = elementOrder(z)
                            Order = modifyOrder(Order)
                            everything = "<p><b>Attention - Double Check!</b></p>"

                            everything = everything + fullDoc(file_name, Order)

                            filedata = {
                                'title': title,
                                'folder': 'F-{}-{}'.format(level, index),
                                'reason': 'test',
                                'fx{}'.format(locID): everything
                            }

                            resp = requests.post(
                                'http://{}{}/rest/1/{}/item/'.format(credentials, url,
                                project), data=filedata)

            #just header files - order element is empty
            else:
                # Upload the header file to the folder if the header file exists
                for file in directory:
                    if(file.startswith("header")):
                        file_name = ele.path + "\\" + file
                        document = Document(file_name)
                        properties = document.core_properties
                        title = properties.title
                        if (properties.title == "[Blank Template]"):
                            title = ele.name
                        requirementsFolder(file_name, absPath, document, title, level, index, project, credentials, folderID, url)

            # Go again into the next subfolder
            foldersClass = classList(folders, ele.path)

            # Function calls itself to investigate folders in lower level
            directorySearch(foldersClass, absPath, index, level, locID, prefix, project, credentials, folderID, url)

# Function that uploads requirements files to the website as a separate items in a specific folder
def requirements(file_name, absPath, document, title, level, index, locID, project, credentials, url):

    print(file_name)

    # Variable that is going to contain all important information that has to be uploaded by the end of the function
    description = ""

    # Zipping the file being uploaded
    z = zipfile.ZipFile(file_name)
    all_files = z.namelist()



    # Check for any images - if they exist, they are located in word/media/ folder of zipped word doc
    # If they are present - save them to local directory and send to server
    if any("word/media/image1.png" in s for s in all_files):

        # Saving image from a word document to a local directory
        saveImg(z, absPath)

        # Adding a string that a function sendImgHtml() creates:
        # sends image to the server, retrieves vital information to create a html which displays that picture
        # returns a string that contains a html formatted picture
        description = sendImgHtml(absPath + "\\word\\media\\image1.png", project, credentials, url) + description


    # Check if there are any graphs - same as images just different image format
    # If so save them to .png format and send them to server
    if any("word/media/image1.tif" in s for s in all_files):
        saveTif(z, absPath)
        tifToPng(absPath)

        description = sendImgHtml(absPath + "\\word\\media\\image1.png", project, credentials, url) + description

    # Again special types of images
    if any("word/media/image1.tmp" in s for s in all_files):

        saveTmp(z, absPath)
        tmpToPng(absPath)

        description = sendImgHtml(absPath + "\\word\\media\\image1.png", project, credentials, url) + description

    if (document.tables != []):
        description = description + tableToHtml(file_name)

    description = wordReadText(file_name) + description

    # Check for any equations
    if (eqnCheck(z)):
        description = "<p><b>Attention - Equations present!</b></p>" + description

    if (objectCheck(z)):
        description = "<p><b>Attention - Object present! (e.g. pdf)</b></p>" + description

    if any("word/media/image2.png" in s for s in all_files):
        description = "<p><b>Attention - Multiple Images Present!</b></p>" + description
    if any("word/media/image2.tif" in s for s in all_files):
        description = "<p><b>Attention - Multiple Images Present!</b></p>" + description
    if any("word/media/image2.tmp" in s for s in all_files):
        description = "<p><b>Attention - Multiple Images Present!</b></p>" + description

    # Gathering data in the correct format - JSON
    filedata = {
        "title": title,
        "folder": "F-{}-{}".format(level, index),
        "reason": "test",
        "fx{}".format(locID): description
    }
    # Uploading information from the word document
    resp = requests.post(
        'http://{}{}/rest/1/{}/item'.format(credentials, url, project), data=filedata)

# Similar as function requirements()
# Only difference is that it uploads to a already existing folder - PUT method (not POST)
def requirementsFolder(file_name, absPath, document, title, level, index, project, credentials, folderID, url):
    description = ""

    z = zipfile.ZipFile(file_name)
    all_files = z.namelist()

    print(file_name)

    # Check for any images
    # If they are present - save them to local directory and send to server
    if any("word/media/image1.png" in s for s in all_files):

        # saves the image from a word document to a local directory
        saveImg(z, absPath)

        description = sendImgHtml(absPath + "\\word\\media\\image1.png", project, credentials, url) + description


    # Check if there are any graphs
    # If so save them to .png format and send them to server
    if any("word/media/image1.tif" in s for s in all_files):
        saveTif(z, absPath)
        tifToPng(absPath)

        description = sendImgHtml(absPath + "\\word\\media\\image1.png", project, credentials, url) + description

    if any("word/media/image1.tmp" in s for s in all_files):
        saveTmp(z, absPath)
        tmpToPng(absPath)

        description = sendImgHtml(absPath + "\\word\\media\\image1.png", project, credentials, url) + description

    if (document.tables != []):
        description = description + tableToHtml(file_name)

    description = wordReadText(file_name) + description

    # Check for any equations
    if (eqnCheck(z)):
        description = "<p><b>Attention - Equation(s) present!</b></p>" + description

    if (objectCheck(z)):
        description = "<p><b>Attention - Object present! (e.g. pdf)</b></p>" + description



    if any("word/media/image2.png" in s for s in all_files):
        description = "<p><b>Attention - Multiple Images Present!</b></p>" + description
    if any("word/media/image2.tif" in s for s in all_files):
        description = "<p><b>Attention - Multiple Images Present!</b></p>" + description
    if any("word/media/image2.tmp" in s for s in all_files):
        description = "<p><b>Attention - Multiple Images Present!</b></p>" + description

    filedata = {
        "title": title,
        "reason": "test",
        "fx{}".format(folderID): description
    }
    resp = requests.put(
        'http://{}{}/rest/1/{}/item/F-{}-{}'.format(credentials, url, project, level, index),
        data=filedata)

# Function that upload standardised tests
def tests(file_name, document, title, level, index, locID, project, credentials, url):
    # Getting the table straight away
    tables = document.tables

    # Lists for upload:
    # description - contains description part of the Tracer table
    description = []
    # results - contains expected results part of the Tracer table
    results = []
    # procedure - contains procedure part of the Tracer table
    procedure = []
    # precond - contains elements that are going to be uploaded to pre-conditions part of Matrix test
    precond = []
    # saving inside table - if there is a table within a table (Expected Results part - Tracer table)
    inTable = []
    # saving a table inside a table - procedure part of the Tracer table
    actionTables = []
    # Boolean values that indicate if there is a table in either expected results or procedure part of the Tracer table, respectively
    tableInTable = False
    tableInAction = False

    # going through the table to find right data
    titID = ""

    for row in tables[0].rows:
        for cell in row.cells:

            # Get description part of the table
            if (cell.text == tables[0].rows[0].cells[1].text):
                for paragraph in cell.paragraphs:
                    description.append(paragraph.text)

            # Get preconditions part of the table
            if (cell.text == tables[0].rows[1].cells[1].text):
                for paragraph in cell.paragraphs:
                    precond.append(paragraph)

            # Get text from the expected results part
            if (cell.text == tables[0].rows[2].cells[1].text):
                for paragraph in cell.paragraphs:
                    if(cell.tables != []):
                        inTable=cell.tables
                        tableInTable = True
                    else:
                        tableInTable = False
                    results.append(paragraph.text)

            # Get text from the procedure part
            if (cell.text == tables[0].rows[3].cells[1].text):
                for paragraph in cell.paragraphs:
                    if(cell.tables != []):
                        actionTables = cell.tables
                        tableInAction = True
                    else:
                        tableInAction = False
                    procedure.append(paragraph.text)

    setUp = []

    # Include edge case if there is nothing in the array
    if (precond == []):
        precond.append("")

    # get the right html format of the text
    preconditions = htmlPrecond(precond)

    # CER/C.E.R - confirm expected result, ER - expected result
    # Adding empty string if C.E.R. appears in the same line and space if C.E.R. appears in a new line
    action, no, resOrder = procedureModify(procedure, tableInAction, actionTables)
    no = no - 1

    # Modifying results - detect tables, lists and reformat the list accordingly
    results = resProcess(results, tableInTable, inTable)

    # Adding backslash if there are any special characters
    results = backSlashAdd(results)
    setUp = backSlashAdd(setUp)
    action = backSlashAdd(action)



    if(len(results)==1):
        res1 = results[0]
    else:
        res1 = ""

    desc=""
    # Check if CER number matches with numbers of ER
    if (len(results) == no):
        print("All good :)")
    else:
        print("Action and Expected results numbers differ!!!")
        print(no, len(results))

        desc = "<p><b>Attention - Check Expected Results!</b></p>"
    for ele in description:
        desc += ele

    # normalising both arrays - should be the same length at the end
    action, results = normalise(action, results, resOrder)

    # Filling setUp list to the same length as other 2 arrays
    # All three arrays need to have the same number of elements
    while len(setUp) < len(action):
        setUp.append("")


    # Adding one results in the end if there are no CER detected in the action section
    if(no == 0):
        results.insert(len(results), res1)
        action.insert(len(action),"")
        setUp.insert(len(setUp),"")



    # Getting the third fxN value for JSON format
    steps = '[{}]'.format(Steps(setUp, action, results))


    #Checking for multiple images - different formats (png, tif & tmp)

    z = zipfile.ZipFile(file_name)
    all_files = z.namelist()

    if any("word/media/image2.png" in s for s in all_files):
        desc = "<p><b>Attention - Multiple Images Present!</b></p>" + desc
    elif any("word/media/image2.tif" in s for s in all_files):
        desc = "<p><b>Attention - Multiple Images Present!</b></p>" + desc
    elif any("word/media/image2.tmp" in s for s in all_files):
        desc = "<p><b>Attention - Multiple Images Present!</b></p>" + desc



    # JSON format data that is transferred to the server
    filedata = {
        # add variable names that change for different files
        "title": title,
        # In which folder we are creating a new folder
        "folder": "F-{}-{}".format(level, index),
        "reason": "test",
        "fx{}".format(locID): desc,
        "fx{}".format((locID+2)): preconditions,
        "fx{}".format((locID+3)): steps
    }

    # post method
    resp = requests.post('http://{}{}/rest/1/{}/item/'.format(credentials, url, project), data=filedata)


    #saving id of the file if it needs to be changed
    num = resp.json()["serial"]

    resp = requests.get('http://{}{}/rest/1/{}/item/{}-{}'.format(credentials, url, project, level, num),
                        data={'[histroy]': 1, '[children]': 'yes'})

    #Checking if uploaded file has 3rd element in fieldValList (STEPS of the test)
    #If not reupload it with error message
    try:
        (resp.json()["fieldValList"]["fieldVal"][2])


    except:
        print("Error, steps weren't imported")
        print(file_name)

        # Uploading with an error massage
        data = {
            "title": title,
            "fx{}".format(locID): desc,
            "fx{}".format((locID + 2)): preconditions,
            "fx{}".format((locID+3)): {"testdata":"<b>Attention - error!</b>","action":"","expected":""},
            "reason": "test"
        }

        resp = requests.put('http://{}{}/rest/1/{}/item/{}-{}'.format(credentials, url, project, level, num), data=data)



#Find folders in a specific directory
def getFolders(array, path):
    folderList = []
    for ele in array:
        if(os.path.isdir(path + "\\" + ele)):
            folderList.append(ele)
    return folderList


#Initialize a list of objects
def classList(array, directory):
    objects = []
    for ele in array:
        obj = folder(ele, directory + "\\" + ele)
        objects.append(obj)
    return objects

#Reading folder title from documents.xml file
def readTitle(file):
    tree = ET.parse(file)
    root = tree.getroot()
    attribute = root.attrib
    return attribute["title"]

#Reading folder type from documents.xml file
def readType(file):
    tree = ET.parse(file)
    root = tree.getroot()
    attribute = root.attrib
    return attribute["type"]

#Reading folder prefix from documents.xml file
def redPrefix(file):
    tree = ET.parse(file)
    root = tree.getroot()
    attribute = root.attrib
    #check correct naming
    return attribute["prefix"]

#Reading folders order structure inside some folder from structure.xml file
def structureFolder(file):
    array = []
    mydoc = minidom.parse(file)
    items = mydoc.getElementsByTagName("structure")
    for ele in items:
        for ele2 in ele.childNodes:
            if (str(ele2.nodeName) == "section"):
                array.append(ele2.attributes["name"].nodeValue)
    return array

#Reading main folders order structure from set.xml file
def structureSet(file):
    array = []
    mydoc = minidom.parse(file)
    items = mydoc.getElementsByTagName("set")
    for ele in items:
        for ele2 in ele.childNodes:
            if (str(ele2.nodeName) == "document"):
                array.append(ele2.attributes["id"].nodeValue)
    return array

#Reading files ordeer structure inside a folder from structure.xml file
def structureXML(file):
    array = []
    mydoc = minidom.parse(file)
    items = mydoc.getElementsByTagName("structure")
    for ele in items:
        for ele2 in ele.childNodes:
            if(str(ele2.nodeName)=="item"):
                array.append(ele2.attributes["id"].nodeValue)
    return array

#reading text from a word document
def wordReadText(file):
    # Opening a word document with python-docx library
    doc = Document(file)

    # List that is going to store information from a word document
    fullText = []
    # List that will be made of smaller (2 element) lists, each smaller list contains text of a paragraph and its style
    arr = []

    # Looping through paragraphs and adding their style to a list called arr
    for para in doc.paragraphs:
        tinyarr = [para.text, para.style.name]

        arr.append(tinyarr)

    # Using list arr to create appropriate formatting of different paragrpahs
    # Only does bullet points - no way to tell when something is a bullet point or a numbered list
    for ele in arr:
        if (ele[1] == "List Paragraph"):
            input = "<ul><li>" + ele[0] + "</li></ul>"
            fullText.append(input)
        else:
            input = "<p>" + ele[0] + "</p>"
            fullText.append(input)

    # Return html formatted text as one string - join different elements of a list together
    # with a new line character "\n"
    return "\n".join(fullText)


#Getting name label from costum word property
#Opens word documents but it doesn't close them properly - try to fix that
def costumProperty(file):
    zipped = zipfile.ZipFile(file)
    csp = ""

    # In case some command fails we catch that exception
    try:
        tree = ET.parse(zipped.open("docProps/custom.xml"))
        root = tree.getroot()
        for child in root:
            if ((child.attrib["name"] == "ID")):
                csp=(child[0].text)
                return str(csp)
        csp = "ZZZZ"
        return str(csp)

    except Exception as e:
        print ('\n\n', e)
        return "!!!"

# Function that finds element id on Matrix website
# Important because id's change for each project
def getMatrixID(name, credentials, project, url):

    resp = requests.get('http://{}{}/rest/1/{}/cat'.format(credentials, url, project))

    # Extricting important information form a large JSON file - need to fine ID's for differnet levels
    for ele in resp.json()['categoryList']['categoryExtended']:
        if(ele['category']['shortLabel']==name):
            ID = ele['fieldList']['field'][0]['id']
            return ID


#saving an image from .docx to a directory
def saveImg(zipped, absPath):
    # opening and reading the image from a zipped file
    image = zipped.open("word/media/image1.png").read()
    f = open("image1.png", "wb")
    f.write(image)

    # Extracting the image to a local directory
    zipped.extract("word/media/image1.png", absPath)

# Saving tif files to a local directory from .docx document
# Same as a saveImg() function with different extension
def saveTif(zipped, absPath):
    image = zipped.open("word/media/image1.tif").read()

    f = open("image1.tif", "wb")
    f.write(image)
    zipped.extract("word/media/image1.tif", absPath)

# Converting tif format to png - using PIL library
def tifToPng(absPath):
    outfile = os.path.splitext(absPath + "\\word\\media\\image1.tif")[0] + ".png"
    try:
        im = Image.open(absPath + "image1.tif")
        im.thumbnail(im.size)
        im.save(outfile, "PNG", quality=100)
    except Exception as e:
        print(e)

# renaming .tmp file to png files
def saveTmp(zipped, absPath):
    image = zipped.open("word/media/image1.tmp").read()
    f = open("image1.tmp", "wb")
    f.write(image)
    zipped.extract("word/media/image1.tmp", absPath)

# converting tmp format to png using PIL library
def tmpToPng(absPath):
    outfile = os.path.splitext(absPath + "\\word\\media\\image1.tmp")[0] + ".png"
    try:
        im = Image.open(absPath + "\\word\\media\\image1.tmp")
        im.thumbnail(im.size)
        im.save(outfile, "PNG", quality=100)
    except Exception as e:
        print(e)

#send image to the server and returns fileId and key of the image on the server
def sendImgHtml(file, project, credentials, url):
    # Getting the correct JSON format  for picture upload
    files = {
        'upload file': open(file, 'rb')
    }

    # Uploading the picture
    resp = requests.post('http://{}{}/rest/1/{}/file'.format(credentials, url, project), files=files)

    # Creating a string with the correct key and fileId that the are returned by the server when an image is uploaded
    text = '<img src="http://{}/rest/1/{}/file/{}?key={}"><p><br></p>'.format(url, project,
        resp.json()['fileId'], resp.json()['key'])

    # Returning the html formatted string that displays uploaded image
    return text

# Equation that checks if there are any equations present in the .docx file
# if they are present return True else return False
def eqnCheck(zipped):

    doc = zipped.open("word/document.xml").read()

    # "m:oMathPara" and "m:oMath" are tags in a .xml file that indicate there is an equation present
    if("m:oMathPara" in str(doc) or "m:oMath" in str(doc)):
        boolVal = True
    else:
        boolVal = False
    return boolVal

# Function that checks if there are any objects present - e.g. pdf
# if the object is present it returns true, else it returns false
def objectCheck(zipped):
    doc = zipped.open("word/document.xml").read()

    # "w:object" is a tag in a .xml file that indicates an object is present in a word document
    if ("w:object" in str(doc)):
        boolVal = True
    else:
        boolVal = False
    return boolVal

# Function that converts a word table into a html formatted language
def tableToHtml(file):

    # initializing yattag library
    doc, tag, text = Doc().tagtext()

    document = Document(file)
    tables = document.tables

    # with tag() - uses yattag library
    # All it does: adds html tags before at the text so that a table is created

    # Looping through list of tables
    for table in tables:
        with tag("div"):
            with tag("table", klass="table table-bordered"):
                with tag("tbody"):
                    for row in table.rows:
                        with tag("tr"):
                            for cell in row.cells:
                                with tag("td"):
                                    # This condition probably needs to be changed - check different styles
                                    # if (cell.paragraphs[0].style.name == "List Paragraph"): - doesn't work all the time
                                    if (len(cell.paragraphs) > 1):
                                        with tag("ol", type="1"):
                                            array = []
                                            count = 0
                                            temp = []
                                            for paragraph in cell.paragraphs:
                                                temp.append(paragraph.text)

                                            for ele in temp:
                                                count = count + 1

                                                if (ele.endswith(":") or ele.endswith(": ")):
                                                    # Creating a small array of elements that should be grouped
                                                    # They are inserted together to the bigger array as one element (list)
                                                    tinyarr = []
                                                    tinyarr.append(ele)
                                                    no = count

                                                    while True:
                                                        if (no < len(temp)):
                                                            tinyarr.append(temp[no])
                                                            if (temp[no].endswith(".")):
                                                                temp.remove(temp[no])
                                                                array.append(tinyarr)
                                                                break

                                                            temp.remove(temp[no])
                                                        else:
                                                            break
                                                else:
                                                    #temporary one element list:
                                                    temp1 = [ele]

                                                    array.append(temp1)
                                            results = []

                                            for ele in array:


                                                if (ele[0].endswith(":") or ele[0].endswith(": ")):
                                                    with tag("li"):
                                                        text(ele[0])
                                                        with tag("ol", type="a"):
                                                            # skip the first element
                                                            iterele = iter(ele)
                                                            next(iterele)
                                                            for a in iterele:
                                                                with tag("li"):
                                                                    text(a)
                                                else:
                                                    with tag("li"):
                                                        text(ele[0])

                                    else:
                                        for paragraph in cell.paragraphs:

                                                with tag("p"):
                                                    text(paragraph.text)

    return doc.getvalue()

####################################
# Functions for importing TESTS
####################################

#Getting the correct html formatting for normal texts (does not do styling)
def htmlText(arr):
    doc, tag, text = Doc().tagtext()
    for ele in arr:
        with tag("p"):
            text(ele)
    return doc.getvalue()

# Function that deals with pre-conditions (Matrix) information
def htmlPrecond(paragraphs):
    # initializing yattag library
    doc, tag, text = Doc().tagtext()

    # Creating a numbered list
    with tag("ol", type="1"):
        count=0
        for para in paragraphs:

            # Checks if current paragraph style equals to "List Paragraph"
            if(para.style.name == "List Paragraph"):

                    # Check if that paragraph ends with a colon
                    if (para.text.endswith(":") or para.text.endswith(": ")):

                        # Add second level list - letter list
                        with tag("li"):
                            text(para.text)
                        with tag("ol", type="a"):

                            # Looping through the rest of the elements to see where the list ends
                            for ele2 in paragraphs[(count+1):]:

                                if(ele2.style.name == "List Paragraph"):
                                    with tag("li"):
                                        text(ele2.text)
                                # When list ends break the loop
                                else:
                                    count = count + 1
                                    break
                                count = count + 1

                        break
                    else:
                        with tag("li"):
                            text(para.text)

            # First 2 elements have to be included in a list
            # Sometimes first two elements don't have a Paragraph List style, but are always included in the list
            # This elif accounts for that
            elif(count==0 or count == 1):
                with tag("li"):
                    text(para.text)

            # If there is an element with .zip file, include it in list even if it doesn't have Paragraph List style
            elif(".zip" in para.text):
                with tag("li"):
                    text(para.text)
            # Otherwise break the loop
            else:
                break
            count = count + 1

    # Check if the list reached it's ned
    if (count == (len(paragraphs)-1)):
        return doc.getvalue()
    # If not add the rest of the data
    else:
        idx = 0
        for para in paragraphs[count:]:
            # Adding normal text as a spearate paragraphs
            if(para.style.name == "Normal"):
                with tag("p"):
                    text(para.text)
            # If there is another list add it as a bullet point
            else:
                if(para.text.endswith(":") or para.text.endswith(": ")):
                    with tag("p"):
                        text(para.text)
                    with tag("ul"):
                        for ele2 in paragraphs[(count+idx+1):]:
                            with tag("li"):
                                text(ele2.text)
            idx = idx + 1

        return doc.getvalue()

#Function that modifies procedure array to add spaces & null character to the correct place corresponding to the check for results
#Also deletes CER phrase so that it doesnt appear on Matrix (CER - Confirm Expected Results)
def procedureModify(arr, boolTable, actionTables):
    count = 1
    no = 0

    action = []

    # List that stores all CER that already happened
    history = []

    # List that stores order of CER
    resNum = []

    tableCount = 0

    # Boolean value that is True when there are lists and tables present
    list_table = False
    # Boolean value that is true when there are lists present
    listPresent = False
    for ele in arr:
        #Possible to add other cases where the condition holds true.
        checkdata = "confirm expected result {}".format(count)
        lowerWord = ele.lower()
        match = SequenceMatcher(None, lowerWord, checkdata).find_longest_match(0, len(lowerWord), 0, len(checkdata))

        #removing any spaces at the end of the sentence
        while(lowerWord.endswith(" ")):
            lowerWord = lowerWord[0:len(lowerWord)-1]

        if(lowerWord.endswith(".")):
            lowerWord = lowerWord[0:len(lowerWord)-1]

        # not used at the moment - idea: use it to get how accurately the parts of the strings match
        # if the accuracy is about a certain threshold count it as if they are the same
        if (match.size != 0):
            newSeq = (lowerWord[match.a: match.a+match.size])
            accuracy = SequenceMatcher(None, newSeq, checkdata).ratio()

        else:
            accuracy = 0


        if(lowerWord.endswith("confirm expected result {}".format(count))):

            resNum.append(count-1)
            count += 1
            if ((lowerWord.startswith("confirm expected result {}".format(count-1)) != True)):
                #removing phrase: Confirm Expected Result [].
                #if(no>0):

                # Double check that this gives correct string
                history.append(lowerWord[-26:])

                ele = ele[:-27]

                action.append("*")
                action.append(ele)

            else:

                history.append(lowerWord)
                action.append("+")

        # Might need more testing - threshold value needs to be adjusted (0.90 means 90% matching - needs testing if it's correct)
        elif(accuracy>0.90):




            # Finding the same element
            # For now only works if the CER statement is separate, i.e. is not at the end of another statement
            if any(x == lowerWord for x in history):

                idx = 0
                for ele2 in history:

                    if(lowerWord == ele2):


                        resNum.append(idx)
                    idx += 1
                action.append("+")
            else:
                count += 1
                try:
                    # Finding a number of the CER - last element of the CER string
                    # deducting by one because python list are 0-based
                    no = int(lowerWord[-1]) - 1
                    resNum.append(no)

                # Catching any unanticipated exceptions
                except:
                    resNum.append(count-2)

                # Check if the statment does not start with CER
                if ((lowerWord.startswith("confirm expected result {}.".format(count - 1)) != True) and
                        (lowerWord.startswith("confirm expected result {}".format(count - 1)) != True)):
                    # If CER is at the end and not beginning add "*" indicator
                    action.append("*")
                    action.append(ele)
                    # Add that CER to history
                    history.append(lowerWord)
                else:
                    # statement starts with CER append "+" indicator
                    action.append("+")
                    # Add that CER to history
                    history.append(lowerWord)

        else:

            # When there is no table
            if(boolTable==False):

                #####################################################
                # Possibly to be added                              #
                # Checking for multi leveled lists                  #
                #####################################################

                if(ele.endswith(":") or ele.endswith(": ")):
                    listPresent=True

                # Just add element to action list normally - no special formatting
                action.append(ele)

            #Accounting for tables - section that is not Confirm Expected Result
            else:

                doc, tag, text = Doc().tagtext()
                if(tableCount!=len(actionTables)):
                    if (ele.endswith(":") or ele.endswith(": ")):

                        with tag("p"):
                            text(ele)
                        with tag("div"):
                            with tag("table", klass="table table-bordered"):
                                with tag("tbody"):
                                    for row in actionTables[tableCount].rows:
                                        with tag("tr"):
                                            for cell in row.cells:
                                                with tag("td"):
                                                    for paragraph in cell.paragraphs:
                                                        with tag("p"):
                                                            text(paragraph.text)
                        tableCount += 1

                    else:
                        if (ele != ""):
                            with tag("p"):
                                text(ele)
                    action.append(doc.getvalue())
                else:
                    action.append("<p>" + ele + "</p>")
                    list_table = True
        no = no + 1
    # Part that checks if the number of tables transferred with the total number of tables
    # If they differ add a warning sign that tells user how many tables are missing
    if(tableCount!=(len(actionTables))):
        miss = "<p><b>Attention - {} table(s) missing!</b></p>".format((len(actionTables)-tableCount))
        action.append(miss)
    if(list_table == True):
        miss = "<p><b>Attention - list(s) and table(s) present!</b></p>"
        action.append(miss)

    if(listPresent == True):
        miss = "<p><b>Attention - list(s) present!</b></p>"
        action.append(miss)

    # returning action list, count - number of CER detected and
    # resNum - list with integers that indicate order of CER as they appear in Procedure row of Tracer test table
    return action, count, resNum

#Function that groups elements together if there different level lists
def resProcess(arr, boolTable, inTable):
    count = 0
    array = []


    # Boolean variable that checks whether a full stop is included at the end of the text
    # If not it account for maximum of two separate lists (checks for ":" character)
    # If there is only one ":" character it adds all other text
    fullStopPresent = False


    # Just check for list levels if there are no tables present
    if (boolTable == False):

        #grouping one upper level element with lower level elements
        for ele in arr:
            count = count + 1

            if (ele.endswith(":") or ele.endswith(": ")):
                #Creating a small array of elements that should be grouped
                #They are inserted together to the bigger array as one element (list)
                tinyarr = []
                tinyarr.append(ele)
                no = count

                # tempArr - list for saving data when there is list present
                # Later gets appended to the main array with all other data
                tempArr = []

                # Python version of a do-while loop
                while True:
                    # Statement to ensures no doesn't go out of index when looking at list arr
                    if(no<len(arr)):
                        tinyarr.append(arr[no])

                        # Condition for exiting the loop
                        if(arr[no].endswith(".") or arr[no].endswith(". ")):
                            arr.remove(arr[no])
                            tempArr.append(tinyarr)
                            fullStopPresent == True
                            break
                        # necessary condition for exiting the loop - colon indicates next list
                        elif(arr[no].endswith(":") or arr[no].endswith(": ")):
                            #arr.remove(arr[no])
                            tinyarr.pop()
                            tempArr.append(tinyarr)

                            break

                        arr.remove(arr[no])

                    else:

                        # When there is no full stop or second ":" character, tempArr is an empty list
                        # Add all elements to the array in a way they are all group in a lower level
                        # Except the first element with ":" character
                        if(tempArr==[]):
                            tempArr.append(tinyarr)

                        # If the second ":" character was detected add second part as a separate list
                        # Represent separate list
                        elif(fullStopPresent!=True):
                            tempArr.append(tinyarr)

                        break
                # Unpacking of a list
                [tempArr]=tempArr
                array.append(tempArr)
            else:
                abc = [ele]
                array.append(abc)

        #Creating appropriate html formatting
            results = []

            for ele in array:

                doc, tag, text = Doc().tagtext()
                if (ele[0].endswith(":") or ele[0].endswith(": ")):
                    text(ele[0])
                    with tag("ol", type="a"):
                        #skip the first element - one with colon has already been added 2 lines above
                        iterele = iter(ele)
                        next(iterele)
                        for a in iterele:
                            if (a.endswith(":") or a.endswith(": ")):
                                msg = "<p><b>Attention - More than 2 list levels!</b></p>"
                                text(msg)
                            else:
                                with tag("li"):
                                    text(a)
                else:
                    with tag("p"):
                        text(ele[0])
                results.append(doc.getvalue())

    #For table within a table
    else:
        results = []

        # Check if there is only one table
        if(len(inTable)==1):

            #Check if you need to add some text here


            #boolean value that is true if a paragraph ends with ":"
            colbool = False


            # Adds one table
            # Needs testing!! (Works for action part)
            for ele in arr:
                doc, tag, text = Doc().tagtext()
                if(ele.endswith(":") or ele.endswith(": ")):
                    colbool = True

                    with tag("p"):
                        text(ele)
                    with tag("div"):
                        with tag("table", klass="table table-bordered"):
                            with tag("tbody"):
                                for row in inTable[0].rows:
                                    with tag("tr"):
                                        for cell in row.cells:
                                            with tag("td"):
                                                for paragraph in cell.paragraphs:
                                                    with tag("p"):
                                                        text(paragraph.text)

                else:
                    if(ele != ""):
                        with tag("p"):
                            text(ele)
                results.append(doc.getvalue())

            #If no ":" detected group all paragraphs together and paste table at the bottom
            if (colbool==False):
                doc2, tag2, text2 = Doc().tagtext()


                with tag2("div"):
                    with tag2("table", klass="table table-bordered"):
                        with tag2("tbody"):
                            for row in inTable[0].rows:
                                with tag2("tr"):
                                    for cell in row.cells:
                                        with tag2("td"):
                                            for paragraph in cell.paragraphs:
                                                with tag2("p"):
                                                    text2(paragraph.text)

                results.append(doc2.getvalue())
                results = ["".join(results)]

        # More than one table
        else:

            list_table = False

            tableCount = 0
            for ele in arr:
                combine = ""

                # Safety that tableCount doesn't go out of index (of inTable)
                if(tableCount!=len(inTable)):
                    # Condition when to add a table
                    if (ele.endswith(":") or ele.endswith(": ") or ele.endswith("are.") or ele.endswith("are. ") or ele.endswith("are") or ele.endswith("are ")):
                        # Adding one paragrpah before the table - one with a colon or something similar
                        combine = "<p>" + ele + "<p>"
                        # Making a table
                        doc, tag, text = Doc().tagtext()
                        with tag("div"):
                            with tag("table", klass="table table-bordered"):
                                with tag("tbody"):
                                    for row in inTable[tableCount].rows:
                                        with tag("tr"):
                                            for cell in row.cells:
                                                with tag("td"):
                                                    for paragraph in cell.paragraphs:
                                                        with tag("p"):
                                                            text(paragraph.text)
                        combine = combine + doc.getvalue()
                        results.append(combine)
                        tableCount += 1
                # If there is more ":" characters than tables add the element normally
                # and indicate lists and tables are present at the same time
                else:
                    results.append("<p>" + ele + "</p>")
                    list_table = True

            # If nothing was added to the results list - do some fixing
            if(results==[]):

                # If table lenght doesn't match paragraph length just group paragraphs
                if(len(inTable)!=len(arr)):
                    for ele in arr:
                        results.append(ele)

                # If table length matches paragraph length, group one paragraph with one table
                # Catches any exceptions that might be missed in the loop above
                # (if paragraphs don't end with correct ending)
                else:
                    tableCount=0
                    for ele in arr:
                        combine = "<p>" + ele + "<p>"
                        doc, tag, text = Doc().tagtext()
                        with tag("div"):
                            with tag("table", klass="table table-bordered"):
                                with tag("tbody"):
                                    for row in inTable[tableCount].rows:
                                        with tag("tr"):
                                            for cell in row.cells:
                                                with tag("td"):
                                                    for paragraph in cell.paragraphs:
                                                        with tag("p"):
                                                            text(paragraph.text)
                        combine = combine + doc.getvalue()
                        results.append(combine)
                        tableCount += 1

            # Add a warning of how many tables are missing
            # when total number of tables doesn't match number of tables that were uploaded
            if (tableCount != len(inTable)):
                miss = "<p><b>Attention {} table(s) missing</b></p>".format(len(inTable)-tableCount)
                results.append(miss)
            if(list_table):
                miss = "Attention - list(s) and table(s) present!"
                results.append(miss)
    # Return results list back
    return results




#Advance backslash adder - atm checks for " and \ characters, add if there are others
def backSlashAdd(arr):


    for count, ele in enumerate(arr):
        #variable that counts how many of corresponding characters are present in one element of the list
        no = 0
        #go through every character in an element
        for idx, char in enumerate(ele):
            #check for \
            if(char=="\""):
                #variable know how many characters are added
                l = 1
                #if its the first character that fits the definition save it first index
                if(no == 0):
                    check = idx
                else:
                    check = 0
                #gives us distance from the first element
                sub = idx - check

                #checks if its the first element
                if(sub != idx):
                    ele = ele[:idx] + "\\" + ele[idx:]
                #if not we adjust the position of insertion with no and l variables
                #was a problem - when we change one element the index stays the same until next loop
                #when an element is added we need to adjust the indexing for elements that are inserted later
                else:
                    ele= ele[:idx+no*l] + "\\" + ele[idx+no*l:]
                no = no + 1
                arr.pop(count)
                arr.insert(count, ele)



            #same as above just different character

            elif (char == "\\"):
                l = 1
                if (no == 0):
                    check = idx
                else:
                    check = 0
                sub = idx - check

                if (sub != idx):
                    ele = ele[:idx] + "\\" + ele[idx:]
                else:
                    ele = ele[:idx + no*l] + "\\" + ele[(idx + no*l):]
                no = no + 1
                arr.pop(count)
                arr.insert(count, ele)


            # Not exactly sure why needed, because it doesn't change a thing
            # However, without it program is not able to detect tab character (/t)
            # which prevents it from displaying data on the website
            elif (char == "/"):
                l = 1
                if (no == 0):
                    check = idx
                else:
                    check = 0
                sub = idx - check

                if (sub != idx):
                    ele = ele[:idx] + ele[idx:]
                else:
                    ele = ele[:idx + no*l] + ele[(idx + no*l):]
                #no = no + 1
                arr.pop(count)
                arr.insert(count, ele)

            #removing any tab characters
            elif (char == "\t"):

                ele = ele.replace("\t", "")
                arr[count] = ele

            #removin any new line characters
            elif(char == "\n"):

                ele = ele.replace("\n", "")
                arr[count] = ele


    return arr

#Normalising expected results to given actions
def normalise(action, results, resOrder):

    count = 0
    idx = 0

    bool = True
    normRes = []
    normAct = []
    #going through action array (has flags where should be a result or not)
    for ele in action:

        #space present solo CER
        if(ele == "+"):
            try:
                if(resOrder[idx]<len(results)):
                    #adding result element with correct index to a new array

                    #Takes care if CER phrase appears as first thing in an action array
                    if(count == 0):
                        normAct.append("")
                    # Takes care if CER appears in a secodn place in action list (probably could be done with last if statment)
                    # However it's working so I'll keep it
                    if(count==1 and normRes[0]==""):
                        normRes.pop()

                    # Takes care if 2 or more CER phrases appear one after another
                    if(count>0 and (action[count-1]=="+" or action[count-1]=="*") and (action[count]=="+" or action[count]=="*")):
                        normAct.append("")
                    # Takes care if CER appears where ever in a list after first 2 elements and previous element was not CER
                    if(count>1 and normRes[-1]==""):
                        normRes.pop()

                    normRes.append(results[resOrder[idx]])
                    #that index needs to be incremented for next element
                    idx = idx + 1

                    #Boolean value to check when a space should be added for normalise results array
                    bool = True
                else:
                    normRes.append("<p><b>Attention - Result missing!</b></p>")
                    normAct.append("")
                    bool = True
            except Exception as e:
                print(e)
                normRes.append("<p><b>Attention - Result missing!</b></p>")
                normAct.append("")

        #empty string detected in the action array
        elif(ele == "*"):

            try:
                #we add another result
                if(resOrder[idx]<len(results)):

                    normRes.append(results[resOrder[idx]])
                    #increment index
                    idx = idx + 1
                    #set the bool to false because we don't want empty space in the normalised results array
                    bool = False
                else:
                    normRes.append("<p><b>Attention - Result missing2!</b></p>")
                    bool = False
            except:
                normRes.append("<p><b>Attention - Result missing2!</b></p>")
                bool = False


        else:

            #when there is no flag new check the state of the bool to add an empty space or not
            if(bool):
                normRes.append("")
            #we have to add element from action array
            normAct.append(ele)
            bool = True


        count = count + 1

    return normAct, normRes

#Function that returns a list of dicitonaries for json file third fxN element
def Steps(setup, action, results):

    s = ""
    for set, act, res in zip(setup, action, results):
        individual = '{"testdata":"' + set + '","action":"' + act + '","expected":"' + res + '"}'
        if (s != ""):
            s = s + "," + str(individual)
        else:
            s += (individual)
    return s


########################################################################################################################
# Function that takes a zipped file and looks through its document.xml file where it finds order of different elements #
########################################################################################################################
#Probably needs some modifications - hard to predict every possible scenario
#############################################################################

# Function that opens document.xml from a zipped word document and reads information from it's tags
def elementOrder(zipped):
    mydoc = minidom.parse(zipped.open("word/document.xml"))

    items = mydoc.getElementsByTagName('w:body')

    elArray = []
    for elem in items: #Highest level - body

        for ele2 in elem.childNodes: #Secodn level - table&p

            #Checks if there are any empty spaces - level 2
            if(str(ele2.firstChild) != "None" and str(ele2.nodeName)!="w:sectPr"):

                # tag that represents a table
                if(str(ele2.nodeName)=="w:tbl"):
                    tinyArr = []
                    tinyArr.append("Table")

                    row = 0
                    for ele3 in ele2.childNodes:
                        col = 0

                        #Table within a table
                        if(str(ele3.nodeName)=="w:tr"):
                            row = row + 1

                            for ele4 in ele3.childNodes:
                                col = col + 1
                                for ele5 in ele4.childNodes:
                                    if(str(ele5.nodeName)=="w:tbl"):
                                        tinyArr.append("table")
                                        #index where the sub table is - e.g. 2x3 - 2nd row, 3rd column
                                        tinyArr.append("{}x{}".format(row, col))

                    tinyArr.append("{}x{}".format(row,col))
                    elArray.append(tinyArr)
                else:
                    #detection lists
                    if(str(ele2.firstChild.nodeName) == "w:pPr"):
                        #elArray.append("ListParagraph")
                        tinyArr = []
                        ######################
                        #Add differnet styles
                        ######################

                        tinyArr.append("ListParagraph")

                        # Safety - handling exceptions
                        try:
                            # Finding type of the list - ordered/unordered/bulletpoints
                            #print(ele2.firstChild.nodeName)
                            if ((ele2.firstChild.childNodes[1].lastChild.attributes["w:val"].nodeValue) == "1"):
                                tinyArr.append("Number")

                            #Check all numbers again before using it
                            elif ((ele2.firstChild.childNodes[1].lastChild.attributes["w:val"].nodeValue) == "2" or
                                  (ele2.firstChild.childNodes[1].lastChild.attributes["w:val"].nodeValue) == "14" or
                                    (ele2.firstChild.childNodes[1].lastChild.attributes["w:val"].nodeValue) == "15"):
                                tinyArr.append("Bullet")

                            else:
                                tinyArr.append("Heading")

                            #finding level of the list - up to 3 levels
                            if((ele2.firstChild.childNodes[1].firstChild.attributes["w:val"].nodeValue)=="0"):
                                # 0 - first level
                                tinyArr.append("0")

                            elif ((ele2.firstChild.childNodes[1].firstChild.attributes["w:val"].nodeValue)=="1"):
                                # 1 - second level
                                tinyArr.append("1")

                            else:
                                # 2 - third level
                                tinyArr.append("2")

                            elArray.append(tinyArr)
                        except Exception as e:
                            print(e)

                    #Finding equations - Hopefully this covers all cases - only 2 elements checked (first and last)
                    elif(str(ele2.firstChild.nodeName) == "m:oMathPara"):
                        elArray.append(["Equation"])

                    elif(str(ele2.lastChild.nodeName)== "m:oMathPara"):
                        elArray.append(["Equation"])


                    #detecting picture:
                    else:
                        if(str(ele2.childNodes[0].firstChild)!="None"):
                            if(str(ele2.childNodes[0].firstChild.nodeName)=="w:rPr"):
                                elArray.append(["Picture"])
                            else:
                                elArray.append(["Paragraph"])
                        else:
                           elArray.append(["Paragraph"])

    # Returning a list of elements that should be in order as they appear in a word document
    return elArray

#Paragraph reader for non-generic files
def paragraphRead(file_name, no):

    #reading one paragraph at a time - depends on the index given in the argument
    doc = Document(file_name)

    # Safety if no variable exceeedes doc.paragraph max index
    try:
        paragraph = doc.paragraphs[no]
        return str(paragraph.text)
    except Exception as e:
        print(e)
        # Trying to handle if doc.paragraphs doesn't have a lot of elements
        # Probably doesn't work properly - can be changed (or even deleted)
        if(no == 1 or no == 2):
            return ""
        else:
            #Try to find a solution to return error when file is empty
            return " "

#Table creator for non-generic files
def tableHtml(file, no, row1=0, col1=0):

    doc, tag, text = Doc().tagtext()

    document = Document(file)

    # Finding a certain table - which one depends on no
    table = document.tables[no]

    # Initializing table in html format which can include table list and table inside of it
    with tag("div"):
        with tag("table", klass="table table-bordered"):
            with tag("tbody"):
                for row in table.rows:
                    with tag("tr"):
                        for cell in row.cells:
                            with tag("td"):
                                # This condition can probably be improved - check different styles
                                # At the moment just checks if there are multiple paragraphs inside the same cell
                                if (len(cell.paragraphs) > 1):
                                    with tag("ol", type="1"):
                                        for paragraph in cell.paragraphs:
                                            with tag("li"):
                                                text(paragraph.text)
                                else:
                                    for paragraph in cell.paragraphs:
                                        #check if there is a table present in a table and adds it in
                                        try:
                                            if(table.rows[row1-1].cells[col1-1].tables != [] and cell.tables != []):
                                                # When table is present add html formatted table to that cell
                                                with tag("p"):
                                                    inTable = cell.tables[0]
                                                    with tag("table", klass="table table-bordered"):

                                                        with tag("tbody"):
                                                            for inRow in inTable.rows:
                                                                with tag("tr"):
                                                                    for inCell in inRow.cells:
                                                                        with tag("td"):
                                                                            for inParagraph in inCell.paragraphs:
                                                                                with tag("p"):
                                                                                    text(inParagraph.text)
                                            else:
                                                with tag("p"):
                                                    text(paragraph.text)

                                        # Exceptions happen when Tables are strange format - goes out of max list index
                                        except:

                                            print("!Strange Table format!")
    # Return html format of a table
    return doc.getvalue()

# Modifying order a bit so it's easier to find where a list starts
# Needs some work
def modifyOrder(array):
    modified = []
    for idx, ele in enumerate(array):
        # If the element is first in a list add number 1 to the end of
        # changes "ListParagraphs" to "ListParagraphs1
        if(ele[0]=="ListParagraph" and ele[1]!="Heading"):

            if(array[idx-1][0]!="ListParagraph" and array[idx-1][0]!="ListParagraph1"):
                current = ele
                current[0]="ListParagraph1"
                modified.append(current)

            else:
                modified.append(ele)
        else:
            modified.append(ele)
    return array

# Function that returns whole document in html form
# Tries it's best, but I feel like it needs a lot of work
# Might be easier to write it from scratch - with elementOrder() function
def fullDoc(file, orderedArray):
    count = 0
    # variable that counts number of paragraphs already in html format
    paraNum = 0
    # variable that counts number of tables already in html format
    tableNum = 0

    # string that will contain a whole html format of a word document
    description = ""

    #Go through each element of the array
    for ele in orderedArray:

        #Check for every individual element - process of elimination
        if(ele[0]=="Paragraph"):

            #Add correct paragraph to the description
            para = paragraphRead(file, paraNum)
            while (para == ""):
                paraNum = paraNum + 1
                para = paragraphRead(file, paraNum)
            para = "<p>" + para + "</p>"

            description = description + para
            paraNum = paraNum + 1

        #Check if there are any tables and tables within them
        elif(ele[0]=="Table"):
            #Modify for
            if(ele[1]=="table"):
                r = int(ele[2][0])
                c = int(ele[2][2])

                description = description + tableHtml(file, tableNum, r, c)
                tableNum = tableNum + 1
            else:
                description = description + tableHtml(file, tableNum)
                tableNum = tableNum + 1

        #Check for any headers - not sure if ordered list is working correctly for Headings - Approach with care
        elif(ele[0]=="ListParagraph"):
            if (ele[1] == "Heading"):
                para = paragraphRead(file, paraNum)
                while (para == ""):
                    paraNum = paraNum + 1
                    para = paragraphRead(file, paraNum)

                para = "<p><b>" + para + "</b></p>"
                description = description + para
                paraNum = paraNum + 1

        #Check if there are any equations (only checks main element - if they are not apart something else)
        elif(ele[0]=="Equation"):
            description = description + "<p><b>EQUATION</b><p>"

        # Ordered/unordered lists (e.g. bullet points) - a bit confusing
        # Tries to create a html formatted list according to the order of elements
        elif(ele[0]=="ListParagraph1" and ele[1]!="Heading"):

            #First check if the list is ordered
            if(ele[1]=="Number"):

                no = 0
                #add the correct tag
                # String that contains ordered (numbers) lsit
                numbering = '<ol type="1">'
                #go through the elements of the list
                for idx, ele2 in enumerate(orderedArray[count:]):


                    #exit the loop if the element is not a part of a list
                    if(ele2[0]!="ListParagraph" and ele2[0]!="ListParagraph1"):

                        break
                    else:
                        #check the level of the element (0 = highest, 1 = second higest ...)
                        if(ele2[2]=="0"):
                            para = paragraphRead(file, paraNum)
                            #Ignore empty character in paragraphs
                            while(para == ""):
                                paraNum = paraNum + 1
                                para = paragraphRead(file, paraNum)

                            #when a character is not empty add it to the list
                            para ="<li>" + para + "</li>"
                            numbering = numbering + para
                            #increment number of paragraph that was read
                            paraNum = paraNum + 1
                            #increment the number of list (used for sublists)
                            no = no + 1
                        #second level:
                        else:
                            letters = '<ol type="a">'
                            #same as above only for one level lower that has letter numbering
                            for ele3 in orderedArray[(count + no):]:
                                if (ele3[0] != "ListParagraph" and ele3[0] != "ListParagraph1"):
                                    break

                                elif(ele3[2]=="1"):

                                    let = paragraphRead(file, paraNum)
                                    while (para == ""):
                                        paraNum = paraNum + 1
                                        let = paragraphRead(file, paraNum)

                                    let = "<li>" + let + "</li>"
                                    letters = letters + let
                                    paraNum = paraNum + 1

                            #closing tags and adding to the while element
                            letters = letters + "</ol>"
                            numbering = numbering + letters
                numbering = numbering + "</ol>"

                description = description + numbering

            # Same as above only for bullet points (unordered lists)
            elif (ele[1] == "Bullet"):

                no = 0

                # String that combines bullet point together
                # Bullet point tag
                numbering = '<ul>'

                for idx, ele2 in enumerate(orderedArray[count:]):


                    if (ele2[0] != "ListParagraph" and ele2[0] != "ListParagraph1"):
                        break

                    else:
                        if (ele2[2] == "0"):
                            para = paragraphRead(file, paraNum)
                            while (para == ""):
                                paraNum = paraNum + 1
                                para = paragraphRead(file, paraNum)

                            para = "<li>" + para + "</li>"
                            numbering = numbering + para
                            paraNum = paraNum + 1
                            no = no + 1
                        # second level:
                        else:
                            letters = '<ul>'
                            for ele3 in orderedArray[(count + no):]:
                                if (ele3[0] != "ListParagraph" and ele3[0] != "ListParagraph1"):

                                    break
                                elif (ele3[2] == "1"):

                                    let = paragraphRead(file, paraNum)
                                    while (para == ""):
                                        paraNum = paraNum + 1
                                        let = paragraphRead(file, paraNum)

                                    let = "<li>" + let + "</li>"
                                    letters = letters + let
                                    paraNum = paraNum + 1

                            letters = letters + "</ul>"
                            numbering = numbering + letters
                numbering = numbering + "</ul>"

                description = description + numbering
        count = count + 1

    #return whole html format
    return description